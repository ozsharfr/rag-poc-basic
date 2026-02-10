import os
import re
from tkinter.messagebox import QUESTION
import torch
import docx
import chromadb
import requests
from typing import List, Dict, Any
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_classic.retrievers import EnsembleRetriever
from langchain_community.retrievers import BM25Retriever
from langchain_core.documents import Document

from rank_bm25 import BM25Okapi
import logging

# logger
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# -----------------------------
# 0) Utils
# -----------------------------

def is_file_in_db(vectorstore, file_name: str) -> bool:
    """Return True if at least one chunk with metadata file_name exists in the Chroma collection."""
    res = vectorstore._collection.get(
        where={"file_name": os.path.basename(file_name)},
        limit=1
    )
    return len(res.get("ids", [])) > 0


def is_valid_analysis(text: str) -> bool:
    """Heuristic guardrail: detect cases where the 'analysis' is just prompt echo / non-answer."""
    if not text:
        return False

    t = text.strip()
    forbidden = [
        "סכם ב-2 משפטים",
        "טקסט לניתוח",
        "פלט:",
        "תשובה:",
        "ענה בעברית בלבד",
    ]
    if any(m in t for m in forbidden):
        return False

    # סלחני יותר לסיכומים קצרים - מעל 10 תווים
    if len(t) < 10:
        return False

    return True


# -----------------------------
# 1) Document Processing
# -----------------------------

def extract_segments(file_path, max_chunk_size=1000): # Specific to DOCX with headings as Bold+Underline, but can be extended to other formats and heading styles.
    """Load DOCX and split to segments by headings (Bold + Underline) and optional chunk size."""
    doc = docx.Document(file_path)
    file_name = os.path.basename(file_path)
    segments = []
    current_heading = "כללי"
    buffer_text = ""

    for p in doc.paragraphs:
        clean_text = p.text.strip()
        if not clean_text:
            continue
        # In case of bold + underline - we consider it a heading (can be extended with more complex logic if needed)
        is_heading = all(run.bold for run in p.runs if run.text.strip()) and \
                     any(run.underline for run in p.runs if run.text.strip())

        if is_heading:
            if buffer_text:
                segments.extend(_split_to_chunks(buffer_text, current_heading, file_name, max_chunk_size, OVERLAP))
            current_heading = clean_text
            buffer_text = ""
        else:
            buffer_text += clean_text + " "

    if buffer_text:
        segments.extend(_split_to_chunks(buffer_text, current_heading, file_name, max_chunk_size , OVERLAP))

    return segments

# chunks with overlap
def _split_to_chunks(text, heading, file_name, max_size, overlap=200):
    """Split long text to overlapping chunks."""
    text = text.strip()
    chunks = []
    
    step = max_size - overlap
    
    for i in range(0, len(text), step):
        chunk_content = text[i:i + max_size]
        chunks.append({
            "file_name": file_name, 
            "heading": heading, 
            "content": chunk_content
        })
        # avoid small chunks at the end
        if i + max_size >= len(text):
            break
            
    return chunks
# -----------------------------
# 2) Local LLM via Ollama (CPU)
# -----------------------------

def load_ollama_model(model_name: str = "qwen2.5:3b-instruct", base_url: str = "http://localhost:11434"):
    def generate(prompt: str, max_new_tokens: int = 250, temperature: float = 0.2, top_p: float = 0.9) -> str:
        try:
            r = requests.post(
                f"{base_url}/api/generate",
                json={
                    "model": model_name,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": temperature,
                        "top_p": top_p,
                        "num_predict": max_new_tokens,
                    },
                },
                timeout=1000,
            )
            r.raise_for_status()
            data = r.json()
            return (data.get("response") or "").strip()
        except Exception as e:
            raise RuntimeError(f"Ollama generation failed for model='{model_name}': {e}") from e

    return generate


def get_expert_analysis(generate_fn, text, context_limit=2000):
    """פרומפט מפושט: סיכום תמציתי בלבד לשיפור השליפה."""
    context = text[:context_limit]

    prompt = f"""ענה בעברית בלבד.
    סכם את עיקרי הטקסט הבא ב-1-2 משפטים. 
    התמקד במונחי מס מרכזיים.
    טקסט לניתוח:
    {context}

    סיכום:
    """.strip()

    analysis = generate_fn(prompt, max_new_tokens=200, temperature=0.0, top_p=0.9).strip()

    if not is_valid_analysis(analysis):
        return ""

    return analysis

# -----------------------------
# 3) Vector DB (Chroma + E5)
# -----------------------------

def initialize_vector_db(db_path="./tax_db", name="tax_knowledge_base", embed_model_path="C:/models/multilingual-e5-large"):
    embeddings = HuggingFaceEmbeddings(
        model_name=embed_model_path,
        encode_kwargs={"normalize_embeddings": True},
        model_kwargs={"device": "cuda" if torch.cuda.is_available() else "cpu"},
    )

    vectorstore = Chroma(
        persist_directory=db_path,
        collection_name=name,
        embedding_function=embeddings
    )
    return vectorstore


def generate_and_save(vectorstore, segments, llm_generate):
    texts, metadatas, ids = [], [], []

    for i, seg in enumerate(segments):
        summary = get_expert_analysis(llm_generate, seg["content"])
        chunk_id = f"{seg['file_name']}_{i}"

        # Add also title and/or summary
        embedded_text = seg["content"].strip()

        texts.append(embedded_text)
        ids.append(chunk_id)

        metadatas.append({
            "file_name": seg["file_name"],
            "heading": seg["heading"],
            "expert_analysis": summary, # הסיכום נשמר כאן רק לשימוש ה-LLM הסופי
            "raw_content": seg["content"],
            "chunk_id": chunk_id,
        })

    vectorstore.add_texts(texts=texts, metadatas=metadatas, ids=ids)


# -----------------------------
# 4) Hybrid Retrieval (BM25 + Vector)
# -----------------------------



class HybridEnsembleRetriever:
    def __init__(self, bm25_retriever, vector_retriever, k: int = 6, ratio_vector: float = 0.7):
        self.bm25 = bm25_retriever
        self.vec = vector_retriever
        self.k = k
        self.ratio_vector = ratio_vector

    def _call(self, retriever, query: str):
        # תומך גם ב- .invoke וגם ב- get_relevant_documents (למקרה של הבדלי גרסאות)
        if hasattr(retriever, "invoke"):
            return retriever.invoke(query)
        if hasattr(retriever, "get_relevant_documents"):
            return retriever.get_relevant_documents(query)
        raise TypeError(f"Retriever {type(retriever)} has no invoke/get_relevant_documents")

    def invoke(self, query: str):
        bm25_docs = self._call(self.bm25, query) or []
        vec_docs  = self._call(self.vec,  query) or []

        # ניקוד פשוט לפי Reciprocal Rank: 1/(rank+1)
        scores: Dict[str, float] = {}
        docs_by_id: Dict[str, Any] = {}

        def doc_key(doc):
            md = getattr(doc, "metadata", {}) or {}
            return md.get("chunk_id") or md.get("id") or md.get("source") or getattr(doc, "page_content", "")[:80]

        def add(docs, weight):
            for rank, doc in enumerate(docs):
                key = doc_key(doc)
                rr = 1.0 / (rank + 1.0)
                scores[key] = scores.get(key, 0.0) + weight * rr
                docs_by_id[key] = doc

        add(bm25_docs, weight=(1.0 - self.ratio_vector))
        add(vec_docs,  weight=self.ratio_vector)

        ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        return [docs_by_id[k] for k, _ in ranked[: self.k]]

class SimpleBM25Retriever:

    def __init__(self, documents: List[Document], k: int = 6):
        self.documents = documents
        self.k = k
        self._tokenized_corpus = [self._tokenize(d.page_content) for d in documents]
        self._bm25 = BM25Okapi(self._tokenized_corpus)

    def _tokenize(self, text: str) -> List[str]:
        # טוקניזציה פשוטה ובטוחה לעברית/אנגלית/מספרים
        # (אפשר לשפר בהמשך עם normalizations לתחיליות/סיומות וכו')
        return re.findall(r"[א-ת]+|[a-zA-Z]+|\d+", text.lower())

    def invoke(self, query: str) -> List[Document]:
        q = self._tokenize(query)
        scores = self._bm25.get_scores(q)
        top_idx = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[: self.k]
        return [self.documents[i] for i in top_idx]
     
def bm25_text(doc_text: str, meta: dict) -> str:
    title = (meta.get("heading") or "").strip()
    summary = (meta.get("expert_analysis") or "").strip()

    parts = []
    if title:
        # משקל גבוה: חוזרים על הכותרת פעמיים
        parts.append(f"כותרת: {title}\nכותרת: {title}")
    if summary:
        # משקל בינוני: שורה אחת
        parts.append(f"סיכום: {summary}")

    parts.append("---")
    parts.append(doc_text)
    return "\n".join(parts)



def get_custom_retriever(vectorstore, k=6, mode="hybrid", ratio_vector=0.7):
    """
    mode options: 
    'vector' - Cosine Similarity - cleanest , most explainable
    'bm25'   - Only textual - will be most useful with llm that could extract keywords well
    'hybrid' - combining both retrievers 
    """
    
    # 1. הכנת Retriever וקטורי
    v_retriever = vectorstore.as_retriever(search_kwargs={"k": k})
    
    if mode == "vector":
        return v_retriever

    # 2. הכנת BM25 - שליפת המסמכים לצורך בניית האינדקס המילולי
    all_data = vectorstore.get()
    langchain_docs = [
    Document(page_content=bm25_text(doc, meta), metadata=meta)
    for doc, meta in zip(all_data.get('documents', []), all_data.get('metadatas', []))
    ]
    b_retriever = SimpleBM25Retriever(langchain_docs, k=k)

    if mode == "bm25":
        return b_retriever

    # 3. שילוב Hybrid (Ensemble)
    if mode == "hybrid":
        return HybridEnsembleRetriever(
        bm25_retriever=b_retriever,
        vector_retriever=v_retriever,
        k=k,
        ratio_vector=ratio_vector
    )

def generate_final_answer(generate_fn, question, retrieved_docs):
    context_segments = []
    for i, doc in enumerate(retrieved_docs):
        chunk_id = doc.metadata.get('chunk_id') or doc.metadata.get('id') or doc.metadata.get('source')
        header = f"--- מקור {i+1} ({doc.metadata.get('heading', 'ללא כותרת')})"
        if chunk_id:
            header += f" - chunk_id: {chunk_id}"
        header += " ---\n"

        segment = (
            f"{header}"
            f"סיכום מקור: {doc.metadata.get('expert_analysis', '')}\n"
            f"טקסט מלא: {doc.page_content}\n"
        )
        context_segments.append(segment)

    full_context = "\n".join(context_segments).strip()

    prompt = (
        "ענה בעברית בלבד.\n"
        "השתמש אך ורק במידע שמופיע בקונטקסט להלן.\n"
        "אסור להשתמש בידע כללי.\n"
        "אם אין בקונטקסט את המידע המאפשר להשיב — השב: \"לא נמצאה תשובה מפורשת במקורות שבמאגר.\"\n\n"
        f"קונטקסט:\n{full_context}\n\n"
        f"שאלה: {question}\n\n"
        "תשובה:"
    )
    logger.info (prompt)
    logger.info("Generating final answer.")
    answer = generate_fn(prompt, max_new_tokens=800, temperature=0.03, top_p=0.9)
    return (answer or "").strip()

def expand_query(generate_fn, original_query):
            prompt = f"הפוך את השאילתה הבאה לרשימת מילים רלוונטיות לחיפוש משפטי (כולל מילים נרדפות): {original_query}"
            expanded = generate_fn(prompt, max_new_tokens=50)
            return f"{original_query} {expanded}"

def generate_imaginary_answer(generate_fn, original_query):
            prompt = f"""קבל שאלה והצע 3 ניסוחים נוספים קצרים בלבד.
    אל תוסיף מידע חדש. רק ניסוחים חלופיים מקצועיים.

    שאלה: {original_query}


    """
            expanded = generate_fn(prompt, max_new_tokens=50)
            return f"{original_query} {expanded}"

# -----------------------------
# 5) Main
# -----------------------------

E5_PATH = "C:/models/multilingual-e5-large"
DOC_DIR_PATH = "./files_misui"
DB_PATH = "./tax_db_summary_"
DB_NAME = "tax_knowledge_base"
# ---- chunking settings ---
MAX_CHUNK_SIZE = 400
OVERLAP = 200
# --- retrieval settings ---
N_RETRIEVED_DOCS = 10
RETRIEVE_MODE = "vector"#"hybrid"  # options: 'vector', 'bm25', 'hybrid'

MODEL_NAME = "gpt-oss:20b"
EXPAND_QUERY = True
GENERATE_MOCK_ANSWER = False

if __name__ == "__main__":
    

    logger.info("Upload models..")
    llm_generate = load_ollama_model(model_name=MODEL_NAME)

    vectorstore = initialize_vector_db(DB_PATH, name=DB_NAME, embed_model_path=E5_PATH)

    # Indexing
    for file_name in os.listdir(DOC_DIR_PATH):
        full_file_path = os.path.join(DOC_DIR_PATH, file_name)
        if (not file_name.endswith('.docx')) or file_name.startswith('~$'):
            continue
        if is_file_in_db(vectorstore, file_name):
            continue

        logger.info(f"--- processing: {file_name} ---")
        segments = extract_segments(full_file_path, max_chunk_size=MAX_CHUNK_SIZE)
        generate_and_save(vectorstore, segments, llm_generate)

    # Queries
    questions = [
        "איזה מסמך עוסק במיסוי מענק לביצוע פוסט-דוקטורט מחוץ לישראל?",
        "איזה מסמך עוסק בשותפות מוגבלת להשקעה בנכסים פיננסיים (קרן גידור) ובמיסוי השותפים?",
        "איזו החלטת מיסוי עוסקת בהנפקת תעודות סל ותעודות בחסר?",
        "איזו החלטת מיסוי עוסקת בנושא מוסד קבע בישראל בהקשר של מחסן וחלקי חילוף?",
        "אילו החלטות מיסוי נוגעות בבקשת פטור ממס רכישה עבור קרקע חקלאית?",
        "איזו החלטת מיסוי עוסקת במיסוי פיצויי הפקעה, ומה הסעיף שעליו היא מתבססת?",
        "מה ההנחיות הכלליות לתושב חוץ לעניין מס הכנסה?",
        "איזו החלטת מיסוי עוסקת במיסוי פיצויי הפקעה, ומה הסעיף שעליו היא מתבססת?",
        "באיזו החלטת מיסוי נקבע כי כתב הרשאה והתחייבות אינו מהווה עסקה במקרקעין?",
        "איזו החלטת מיסוי עוסקת בנטרול תקופת ההקמה בעסקת BOT, ומה אורך תקופת ההקמה?",
        "איזו החלטת מיסוי עוסקת בהתרת הוצאה לאיתור נכס כניכוי מהשבח, ולאיזה סעיף חוק היא מפנה?"
    ]

    for question in questions[:]: ################### TEST
        logger.info(f"\שאילתה: {question}")
        if EXPAND_QUERY:
            question = expand_query(llm_generate, question)
            logger.info(f"\שאילתה מורחבת: {question}")
        if GENERATE_MOCK_ANSWER:
            question = generate_imaginary_answer(llm_generate, question)
            logger.info(f"\שאילתה מורחבת עם ניסוחים חלופיים: {question}")

        for RETRIEVE_MODE in ["hybrid"]:
            i = 0.7
            logger.info(f"\n--- retreival  {RETRIEVE_MODE} ---")
            retriever = get_custom_retriever(vectorstore,  mode=RETRIEVE_MODE, k=N_RETRIEVED_DOCS, ratio_vector=i)
            retrieved_docs = retriever.invoke(question)
            # Generate final answer based on chunks and query
            final_answer = generate_final_answer(llm_generate, question, retrieved_docs)
            logger.info(f"\n {final_answer}")